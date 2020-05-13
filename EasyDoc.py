from tkinter import*
from tkinter import*
from tkinter import ttk
from tkinter import messagebox
import docx
from docx import *
import os
import re
import shutil
from docx.shared import Cm, Pt
fnt='None 15  '
fnt1='None 15 bold '
bg = 'powder blue'
bgtxt='#ffffff'
fg=           '#000000'
fw =600
fh=500
pad=20
wd=14

class window(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("EasyDoc")
        #self.master.geometry('500x600+0+0')
        #self.master.iconbitmap('ala.ico')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.config(bg='powder blue')
        self.master.resizable(width=False, height=False)
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()

        self.lbltitle = Label(self.frame, text='Easy Doc ', font=('Time New Roman', 50, 'bold'),
                              bg='powderblue', fg='black')
        self.lbltitle.grid(row=0, column=0, columnspan=20,pady=40)
        #============================= Buttons =====================================
        self.btnLogin=Button(self.frame,text='Job certification',width=17,height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.Job_certification_window)
        self.btnLogin.grid(row=3,column=0)
        self.btnReset = Button(self.frame, text='Order mission', width=17,height=4,bd=5,font=('Time New Roman',10, 'bold'),command=self.Order_mission_window)
        self.btnReset.grid(row=3, column=1)
        self.btnExit = Button(self.frame, text='......', width=17,height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window1)
        self.btnExit.grid(row=3, column=2)
        self.btnReset = Button(self.frame, text='.....', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window2)
        self.btnReset.grid(row=4, column=0)
        self.btnExit = Button(self.frame, text='.....', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window3)
        self.btnExit.grid(row=4, column=1)
        self.btnExit = Button(self.frame, text='......', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window4)
        self.btnExit.grid(row=4, column=2)
        self.btnReset = Button(self.frame, text='......', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window5)
        self.btnReset.grid(row=5, column=0)
        self.btnExit = Button(self.frame, text='......', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command= self.window6)
        self.btnExit.grid(row=5, column=1)
        self.btnExit = Button(self.frame, text='Exit', width=17, height=4,bd=5,font=('Time New Roman',10, 'bold'),command=self.master.destroy)
        self.btnExit.grid(row=5, column=2)


    def Job_certification_window(self):
            self.newwindow = Toplevel(self.master)
            self.app = JobCertification_window(self.newwindow)
    def Order_mission_window(self):
            self.newwindow = Toplevel(self.master)
            self.app = ordrMission_window(self.newwindow)
    def window1(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window1(self.newwindow)
    def window2(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window2(self.newwindow)
    def window3(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window3(self.newwindow)
    def window4(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window4(self.newwindow)
    def window5(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window5(self.newwindow)
    def window6(self):
            self.newwindow = Toplevel(self.master)
            self.app = Window6(self.newwindow)
class JobCertification_window(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Job Certification")
        #self.master.geometry('500x600+0+0')
        #self.master.iconbitmap('ala.ico')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        #self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.lbltitle = Label(self.frame, text='Job Certification', font=('Time New Roman', 30, 'bold'), bg='powderblue',
                              fg='black')
        self.lbltitle.grid(row=0, column=0, columnspan=20, pady=30)

        Label(self.frame, text='Full name:', bg=bg,  font=fnt1).grid(row=3, column=0)
        Label(self.frame, text='Date of birth:', bg=bg, font=fnt1).grid(row=3, column=2)
        Label(self.frame, text='Place Of birth:', bg=bg, font=fnt1).grid(row=4, column=0)
        Label(self.frame, text='position :', bg=bg, font=fnt1).grid(row=4, column=2)
        Label(self.frame, text='Date Of Start:', bg=bg, fg=fg, font=fnt1).grid(row=5, column=0)
        Label(self.frame, text='Date Of End:', bg=bg, fg=fg, font=fnt1).grid(row=5, column=2)
        svName = StringVar()
        svdateOfBirth = StringVar()
        svplaceOfBirth = StringVar()
        svdateOfStart = StringVar()
        svdateOfEnd = StringVar()
        svDep = StringVar()
        svdateOfStart = StringVar()
        svdateOfEnd = StringVar()
        txtName = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svName)
        txtdateOfBirth = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svdateOfBirth)
        txtplaceOfBirth = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svplaceOfBirth)
        txtDep = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svDep)
        txtdateOfStart = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svdateOfStart)
        txtdateOfEnd = Entry(self.frame, bg=bgtxt, fg=fg, font=fnt,width=wd, textvariable=svdateOfEnd)
        txtName.grid(row=3, column=1, pady=pad)
        txtdateOfBirth.grid(row=3, column=3, pady=pad)
        txtplaceOfBirth.grid(row=4, column=1, pady=pad)
        txtDep.grid(row=4, column=3, pady=pad)
        txtdateOfStart.grid(row=5, column=1, pady=pad)
        txtdateOfEnd.grid(row=5, column=3, pady=pad)



        def paragraph_search(text_value, new):
            doc = docx.Document(svName.get() + '_job certificate.docx')

            result = False
            para_regex = re.compile(text_value)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    if para_regex.search(paragraph.text):
                        result = True
                        paragraph.text = re.sub(text_value, '  ' + new + '  ', paragraph.text)
                        style = doc.styles['Normal']
                        style.font.name = 'Times New Roman'
                        style.font.size = Pt(12)
                        doc.save(svName.get() + '_job certificate.docx')
            return print(result)

        def copyFile():
            doc = docx.Document('certificate.docx')
            NewFileName = svName.get() + '_job certificate.docx'
            shutil.copy('certificate.docx', NewFileName)
            doc.save('certificate.docx')

        def CreatFile():

            if svName.get().strip() == '':
                messagebox.showinfo('', 'Name is empty')
                txtName.focus()
            elif svdateOfBirth.get().strip() == '':
                messagebox.showinfo('', 'date of birth is empty')
                txtdateOfBirth.focus()
            elif svplaceOfBirth.get().strip() == '':
                messagebox.showinfo('', 'place of birth is empty')
                txtplaceOfBirth.focus()
            elif svDep.get().strip() == '':
                messagebox.showinfo('', 'position is empty')
                txtDep.focus()
            elif svdateOfStart.get().strip() == '':
                messagebox.showinfo('', 'date of start is empty')
                txtdateOfStart.focus()
            elif svdateOfEnd.get().strip() == '':
                messagebox.showinfo('', 'dete of end is empty')
                txtdateOfEnd.focus()
            else:

                copyFile()

                A = ['Name ', 'DateOfBirth ', 'PlaceOfBirth ', 'Dep', 'DateOfStart ', 'DateOfEnd']
                B = [svName.get(), svdateOfBirth.get(), svplaceOfBirth.get(), svDep.get(), svdateOfStart.get(),
                     svdateOfEnd.get()]

                paragraph_search(A[0], B[0])
                paragraph_search(A[1], B[1])
                paragraph_search(A[2], B[2])
                paragraph_search(A[3], B[3])
                paragraph_search(A[4], B[4])
                paragraph_search(A[5], B[5])

                svName.set('')
                svdateOfEnd.set('')
                svdateOfBirth.set('')
                svplaceOfBirth.set('')
                svDep.set('')
                svdateOfEnd.set('')
                svdateOfStart.set('')
                messagebox.showinfo('', 'the file has created')
        self.btnCreat=Button(self.frame,heigh=2,font=fnt1, text='Get The File',command=CreatFile).grid(row=6, column=1,columnspan=2, pady=pad,padx=40)
        self.btnExit=Button(self.frame, text='<= Back', command=self.master.destroy).grid(row=7, column=3,columnspan=2, pady=pad,padx=40)

class ordrMission_window(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Ordr Mission")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()

        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window1(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window2(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window3(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window4(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window5(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)
class Window6(Frame):
    def __init__(self,master=None):
        Frame.__init__(self,master)
        self.master=master
        self.master.title("Easy Doc")
        #self.master.geometry('500x600+0+0')
        x = (self.winfo_screenwidth() - fw) / 2
        y = (self.winfo_screenheight() - fh) / 2 - 60
        self.master.geometry('%dx%d+%d+%d' % (fw, fh, x, y))
        self.master.resizable(width=False, height=False)
        self.master.config(bg='powder blue')
        self.frame = Frame(self.master, bg='powder blue')
        self.frame.pack()
        self.btnExit = Button(self.frame, text='Exit', width=10, height=4, bd=5, font=('Time New Roman', 10, 'bold'),
                              command=self.master.destroy)
        self.btnExit.grid(row=10, column=10)

root=Tk()
app=window(root)
app.mainloop()
